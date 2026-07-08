#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将综合实训报告模板 .docx 转换为 .tex 文件
"""

from docx import Document
import os
import re

INPUT_FILE = r'E:\Zhuanzhuan\docs\2025-综合实训报告模版（计科和AI专用） .docx'
OUTPUT_FILE = r'E:\Zhuanzhuan\docs\综合实训报告模版.tex'
IMAGE_DIR = r'E:\Zhuanzhuan\docs\media'

doc = Document(INPUT_FILE)


def escape_latex(text):
    replacements = [
        ('\\', '\\textbackslash{}'),
        ('{', '\\{'),
        ('}', '\\}'),
        ('$', '\\$'),
        ('&', '\\&'),
        ('#', '\\#'),
        ('^', '\\textasciicircum{}'),
        ('_', '\\_'),
        ('~', '\\textasciitilde{}'),
        ('%', '\\%'),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def is_figure_caption(text):
    t = text.strip()
    return bool(re.match(r'^图\d+\s', t)) or t == '图'


def is_table_caption(text):
    t = text.strip()
    return bool(re.match(r'^表\d+\s', t)) or t == '表'


def get_image_path():
    pic = os.path.join(IMAGE_DIR, 'logo.jpeg')
    if os.path.exists(pic):
        return pic.replace('\\', '/')
    return None


# ========== 收集正文段落（跳过封面部分 0-44） ==========
body_paras = []
for i, p in enumerate(doc.paragraphs):
    if i <= 44:
        continue
    body_paras.append({
        'index': i,
        'text': p.text,
        'style': p.style.name if p.style else 'Normal',
        'align': p.alignment,
    })

# ========== 生成 LaTeX ==========
latex_lines = []

# 文档头部
latex_lines.extend([
    r'\documentclass[12pt,a4paper,oneside]{ctexart}',
    r'\usepackage[top=2.3cm, bottom=2.3cm, left=3.2cm, right=3.2cm]{geometry}',
    r'\usepackage{graphicx}',
    r'\usepackage{array}',
    r'\usepackage{booktabs}',
    r'\usepackage{setspace}',
    r'\usepackage{fancyhdr}',
    r'\usepackage{titlesec}',
    r'\usepackage{tocloft}',
    r'\usepackage{hyperref}',
    '',
    r'\hypersetup{',
    r'    colorlinks=true,',
    r'    linkcolor=black,',
    r'    urlcolor=black,',
    r'}',
    '',
    r'\pagestyle{fancy}',
    r'\fancyhf{}',
    r'\fancyfoot[C]{\thepage}',
    '',
    r'\renewcommand{\cfttoctitlefont}{\hfill\Large\bfseries}',
    r'\renewcommand{\cftaftertoctitle}{\hfill}',
    '',
    r'\titleformat{\section}{\Large\bfseries\centering}{\thesection}{1em}{}',
    r'\titleformat{\subsection}{\large\bfseries}{\thesubsection}{1em}{}',
    r'\titleformat{\subsubsection}{\normalsize\bfseries}{\thesubsubsection}{1em}{}',
    '',
    r'\setstretch{1.5}',
    '',
    r'\newcommand{\covername}[1]{{\fontsize{36pt}{40pt}\selectfont\heiti #1}}',
    r'\newcommand{\coversub}[1]{{\fontsize{18pt}{22pt}\selectfont\heiti #1}}',
    '',
    r'\begin{document}',
    '',
])

# ====== 封面 ======
latex_lines.append('% ========== 封面 ==========')
latex_lines.append(r'\begin{titlepage}')
latex_lines.append(r'\centering')

img_path = get_image_path()
if img_path:
    latex_lines.append(r'\vspace*{2cm}')
    latex_lines.append(r'\includegraphics[width=4cm]{' + img_path + '}')
    latex_lines.append(r'\vspace{1cm}')
else:
    latex_lines.append(r'\vspace*{4cm}')

latex_lines.extend([
    '',
    r'\covername{综合实训报告}',
    r'\\[0.5cm]',
    r'\coversub{具体题目}',
    r'\\[2cm]',
    '',
    r'\begin{tabular}{>{\kaishu\bfseries\large}l>{\large}l}',
    r'  \hline',
    r'  姓名： & \\[0.5cm]',
    r'  学号： & \\[0.5cm]',
    r'  班级： & \\[0.5cm]',
    r'  指导老师： & \\[0.5cm]',
    r'  \hline',
    r'\end{tabular}',
    '',
    r'\vfill',
    r'{\kaishu\large 中国 武汉}',
    r'\\[0.3cm]',
    r'{\kaishu\normalsize 二○二四年七月}',
    r'\\[0.2cm]',
    r'{\kaishu\normalsize 2024.07}',
    '',
    r'\end{titlepage}',
    '',
])

# ====== 目录 ======
latex_lines.append('% ========== 目录 ==========')
latex_lines.append(r'\tableofcontents')
latex_lines.append(r'\newpage')
latex_lines.append('')

# ====== 正文处理 ======
latex_lines.append('% ========== 正文 ==========')

i = 0
while i < len(body_paras):
    p = body_paras[i]
    text = p['text']
    style = p['style']

    # 跳过空段落
    if not text.strip():
        i += 1
        continue

    t = text.strip()

    # Heading 1
    if style == 'Heading 1':
        latex_lines.append(r'\section{' + escape_latex(t) + '}')
        latex_lines.append('')
        i += 1
        continue

    # Heading 2
    if style == 'Heading 2':
        latex_lines.append(r'\subsection{' + escape_latex(t) + '}')
        latex_lines.append('')
        i += 1
        continue

    # 图标题 (如 "图1 xxxxx图")
    if is_figure_caption(t):
        latex_lines.append(r'\begin{figure}[htbp]')
        latex_lines.append(r'\centering')
        latex_lines.append(r'\caption{' + escape_latex(t) + '}')
        latex_lines.append(r'\end{figure}')
        latex_lines.append('')
        i += 1
        continue

    # 表标题 (如 "表1 xxxx表")
    if is_table_caption(t):
        latex_lines.append(r'\begin{table}[htbp]')
        latex_lines.append(r'\centering')
        latex_lines.append(r'\caption{' + escape_latex(t) + '}')
        label_key = re.sub(r'\s+', '', t[:10])
        latex_lines.append(r'\label{' + label_key + '}')
        # 跳过后续的空段落
        i += 1
        while i < len(body_paras) and not body_paras[i]['text'].strip():
            i += 1
        # 不在这里关闭 table，等遇到合适的时机关闭
        # 这里插入测试用例表
        test_table = doc.tables[1]
        latex_lines.append(r'\begin{tabular}{|p{3cm}|p{8cm}|}')
        latex_lines.append(r'\hline')
        for row in test_table.rows:
            c0 = escape_latex(row.cells[0].text.strip().replace('\n', ' '))
            c1 = escape_latex(row.cells[1].text.strip().replace('\n', ' '))
            latex_lines.append(r'  \textbf{' + c0 + '} & ' + c1 + r' \\')
            latex_lines.append(r'  \hline')
        latex_lines.append(r'\end{tabular}')
        latex_lines.append(r'\end{table}')
        latex_lines.append('')
        continue

    # 正文段落
    latex_lines.append(r'\par ' + escape_latex(t))
    i += 1

# ====== 结束 ======
latex_lines.append('')
latex_lines.append(r'\end{document}')

# ========== 写入 ==========
with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
    f.write('\n'.join(latex_lines))

print(f'已生成: {OUTPUT_FILE}')
print(f'共 {len(latex_lines)} 行')
